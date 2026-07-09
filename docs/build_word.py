"""
Tạo file Word tổng hợp toàn bộ tài liệu dự án QuizzGame.
Chạy: python3 docs/build_word.py
"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_heading_color(paragraph, color_rgb):
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(*color_rgb)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_cell_bg(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_inline_formatted_run(para, text):
    """Xử lý **bold**, `code`, *italic* trong 1 đoạn text."""
    # Token pattern
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = para.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x3E)
        elif part.startswith('*') and part.endswith('*'):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)

def strip_md_inline(text):
    """Strip **bold**, `code`, *italic* markers — plain text."""
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    return text

# ─── Main parser ──────────────────────────────────────────────────────────────

def parse_md_to_doc(doc, md_text, h1_color=(0x1A, 0x56, 0xDB), heading_offset=0):
    """
    Chuyển markdown đơn giản thành paragraphs trong doc.
    heading_offset: dịch cấp heading (dùng khi embed section con).
    """
    lines = md_text.split('\n')
    i = 0
    in_code = False
    code_lines = []
    code_lang = ''

    while i < len(lines):
        line = lines[i]

        # ── Code block ──
        if line.startswith('```'):
            if not in_code:
                in_code = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                # End of code block
                in_code = False
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.3)
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                run = p.add_run(code_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
                # light gray background via paragraph shading
                pPr = p._p.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F3F4F6')
                pPr.append(shd)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # ── Horizontal rule ──
        if re.match(r'^---+\s*$', line) or re.match(r'^===+\s*$', line):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── Heading ──
        m = re.match(r'^(#{1,6})\s+(.*)', line)
        if m:
            level = len(m.group(1)) + heading_offset
            text  = strip_md_inline(m.group(2))
            level = min(level, 9)
            style_map = {1: 'Heading 1', 2: 'Heading 2', 3: 'Heading 3',
                         4: 'Heading 4', 5: 'Heading 5', 6: 'Heading 6'}
            style = style_map.get(level, 'Heading 6')
            p = doc.add_heading(text, level=min(level, 6))
            if level == 1:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(*h1_color)
            i += 1
            continue

        # ── Table (detect | ... | ) ──
        if line.startswith('|') and '|' in line[1:]:
            # Collect table rows
            table_lines = []
            while i < len(lines) and lines[i].startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # Filter separator rows
            rows_data = []
            for tl in table_lines:
                if re.match(r'^\|[-:\s|]+\|?\s*$', tl):
                    continue
                cells = [c.strip() for c in tl.strip('|').split('|')]
                rows_data.append(cells)
            if not rows_data:
                continue
            max_cols = max(len(r) for r in rows_data)
            tbl = doc.add_table(rows=len(rows_data), cols=max_cols)
            tbl.style = 'Table Grid'
            for ri, row in enumerate(rows_data):
                for ci, cell_text in enumerate(row):
                    if ci >= max_cols:
                        break
                    cell = tbl.cell(ri, ci)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after  = Pt(1)
                    add_inline_formatted_run(p, strip_md_inline(cell_text))
                    for run in p.runs:
                        run.font.size = Pt(9)
                    if ri == 0:
                        set_cell_bg(cell, 'DBEAFE')
                        for run in p.runs:
                            run.bold = True
            doc.add_paragraph()
            continue

        # ── Blockquote ──
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.35)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            add_inline_formatted_run(p, text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            # left border
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:space'), '4')
            left.set(qn('w:color'), '93C5FD')
            pBdr.append(left)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Unordered list ──
        m_ul = re.match(r'^(\s*)([-*+])\s+(.*)', line)
        if m_ul:
            indent = len(m_ul.group(1)) // 2
            text   = m_ul.group(3)
            style  = 'List Bullet' if indent == 0 else 'List Bullet 2'
            p = doc.add_paragraph(style=style)
            p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.2)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            add_inline_formatted_run(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # ── Ordered list ──
        m_ol = re.match(r'^(\s*)\d+[.)]\s+(.*)', line)
        if m_ol:
            indent = len(m_ol.group(1)) // 2
            text   = m_ol.group(2)
            style  = 'List Number' if indent == 0 else 'List Number 2'
            p = doc.add_paragraph(style=style)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after  = Pt(1)
            add_inline_formatted_run(p, text)
            for run in p.runs:
                run.font.size = Pt(10)
            i += 1
            continue

        # ── Empty line ──
        if line.strip() == '':
            i += 1
            continue

        # ── Normal paragraph ──
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        add_inline_formatted_run(p, line.strip())
        for run in p.runs:
            run.font.size = Pt(10.5)
        i += 1

# ─── Build document ───────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    # ── Cover page ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run('QuizzGame')
    r.font.size  = Pt(36)
    r.bold       = True
    r.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Tài liệu Dự án Tổng hợp')
    r.font.size  = Pt(18)
    r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Ôn thi THPT Quốc gia — Web App')
    r.font.size  = Pt(13)
    r.italic     = True
    r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Ngày: 2026-07-05')
    r.font.size = Pt(11)

    add_horizontal_rule(doc)

    # ── Table of contents (manual) ───────────────────────────────────────────
    doc.add_page_break()
    p = doc.add_heading('Mục lục tài liệu', level=1)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)

    toc_items = [
        ('1', 'Tổng quan dự án (PROJECT_OVERVIEW)'),
        ('2', 'Nhật ký tính năng (FEATURE_LOG)'),
        ('3', 'Test Cases'),
        ('4', 'Code Review Log'),
        ('5', 'Glossary (Thuật ngữ)'),
        ('6', 'Workflow'),
        ('7', 'Hướng dẫn Người dùng (User Guide)'),
        ('8', 'Hướng dẫn Admin (Admin Guide)'),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph(style='List Number')
        add_inline_formatted_run(p, title)

    # ── Documents ────────────────────────────────────────────────────────────

    DOCS_DIR = '/Users/quangnd512/Desktop/claude/quiz_dh/docs'

    sections_config = [
        ('PROJECT_OVERVIEW.md',       'Tổng quan dự án',                 (0x1A, 0x56, 0xDB)),
        ('FEATURE_LOG.md',            'Nhật ký tính năng (Feature Log)', (0x05, 0x72, 0x3A)),
        ('TEST_CASES.md',             'Test Cases',                       (0x7C, 0x3A, 0xED)),
        ('CODE_REVIEW_LOG.md',        'Code Review Log',                  (0xC2, 0x41, 0x0D)),
        ('GLOSSARY.md',               'Glossary — Thuật ngữ',             (0x0E, 0x68, 0x7A)),
        ('WORKFLOW.md',               'Workflow phát triển',               (0x15, 0x50, 0x3B)),
        ('guides/user-guide.md',      'Hướng dẫn Người dùng',            (0x1D, 0x4E, 0xD8)),
        ('guides/admin-guide.md',     'Hướng dẫn Admin',                  (0x9F, 0x1A, 0xAE)),
    ]

    for filename, section_title, color in sections_config:
        filepath = os.path.join(DOCS_DIR, filename)
        if not os.path.exists(filepath):
            print(f'  [SKIP] {filename} — không tìm thấy file')
            continue

        with open(filepath, encoding='utf-8') as f:
            content = f.read()

        doc.add_page_break()

        # Section banner
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(4)
        r = p.add_run(f'  {section_title}  ')
        r.font.size  = Pt(14)
        r.bold       = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*color))
        pPr.append(shd)

        doc.add_paragraph()
        parse_md_to_doc(doc, content, h1_color=color)

        print(f'  [OK] {filename}')

    # ── Save ─────────────────────────────────────────────────────────────────
    out_path = '/Users/quangnd512/Desktop/claude/quiz_dh/docs/QuizzGame_TaiLieu_TongHop.docx'
    doc.save(out_path)
    print(f'\n✅ Đã tạo: {out_path}')
    return out_path

if __name__ == '__main__':
    build_doc()
